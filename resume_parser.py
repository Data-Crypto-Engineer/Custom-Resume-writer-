"""
resume_parser.py
-----------------
Handles extracting plain text from uploaded resume files.

Design:
- One extraction function per file format (extract_pdf, extract_docx).
- A single dispatcher function (extract_resume) decides which extractor
  to call based on the file extension.
- To support a new format later (TXT, HTML, LinkedIn export, etc.), you
  only need to write one new extract_* function and register it in the
  EXTRACTORS dictionary below. No other code needs to change.
"""

from __future__ import annotations

import io
from typing import Callable, Dict

import pdfplumber
from docx import Document


class ResumeExtractionError(Exception):
    """Raised when a resume file cannot be read or parsed."""


def extract_pdf(file) -> str:
    """
    Extract plain text from a PDF file.

    Args:
        file: A file-like object (e.g. from st.file_uploader).

    Returns:
        The extracted text as a single string.
    """
    try:
        text_parts: list[str] = []
        with pdfplumber.open(file) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
        text = "\n".join(text_parts).strip()

        if not text:
            raise ResumeExtractionError(
                "No readable text found in this PDF. It may be a scanned "
                "image without selectable text."
            )
        return text

    except ResumeExtractionError:
        raise
    except Exception as exc:  # noqa: BLE001 - surface any pdf parsing issue clearly
        raise ResumeExtractionError(f"Could not read PDF file: {exc}") from exc


def extract_docx(file) -> str:
    """
    Extract plain text from a DOCX file.

    Args:
        file: A file-like object (e.g. from st.file_uploader).

    Returns:
        The extracted text as a single string.
    """
    try:
        # python-docx expects a file path or a file-like object opened in
        # binary mode. Streamlit's UploadedFile already behaves like one,
        # but we wrap it in BytesIO to be safe with the raw bytes.
        file_bytes = file.read()
        document = Document(io.BytesIO(file_bytes))

        paragraphs = [p.text for p in document.paragraphs if p.text.strip()]

        # Also pull text out of any tables, since resumes sometimes use
        # tables for layout (e.g. skills grids).
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())

        text = "\n".join(paragraphs).strip()

        if not text:
            raise ResumeExtractionError("No readable text found in this DOCX file.")
        return text

    except ResumeExtractionError:
        raise
    except Exception as exc:  # noqa: BLE001 - surface any docx parsing issue clearly
        raise ResumeExtractionError(f"Could not read DOCX file: {exc}") from exc


# Maps a lowercase file extension to the function that extracts its text.
# Adding a new format only requires adding one line here.
EXTRACTORS: Dict[str, Callable] = {
    ".pdf": extract_pdf,
    ".docx": extract_docx,
}


def extract_resume(file, filename: str) -> str:
    """
    Dispatch to the correct extractor based on the file extension.

    Args:
        file: A file-like object to extract text from.
        filename: The original filename, used to determine the format.

    Returns:
        The extracted resume text.

    Raises:
        ResumeExtractionError: If the format is unsupported or extraction fails.
    """
    extension = "." + filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    extractor = EXTRACTORS.get(extension)
    if extractor is None:
        raise ResumeExtractionError(
            f"Unsupported file format '{extension}'. "
            f"Supported formats: {', '.join(EXTRACTORS.keys())}"
        )

    return extractor(file)


def build_master_context(resume_texts: list[tuple[str, str]]) -> str:
    """
    Combine multiple extracted resumes into one labeled master context.

    Args:
        resume_texts: A list of (filename, extracted_text) tuples.

    Returns:
        A single string with each resume clearly separated, e.g.:

            ===== Resume 1 (john_resume.pdf) =====
            ...text...

            ===== Resume 2 (john_old_resume.docx) =====
            ...text...
    """
    sections = []
    for index, (filename, text) in enumerate(resume_texts, start=1):
        sections.append(f"===== Resume {index} ({filename}) =====\n{text}")

    return "\n\n".join(sections)
